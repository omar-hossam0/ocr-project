#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
OCR Runner Script - Uses EasyOCR with Arabic reshaper from model folder
This script processes images and PDFs using the custom Arabic OCR model
"""

import sys
import os
import json
import gc
from pathlib import Path

# Add model directory to Python path
MODEL_DIR = Path(__file__).parent.parent / "model"
sys.path.insert(0, str(MODEL_DIR))

try:
    import easyocr
    import cv2
    import numpy as np
    from PIL import Image
    # Compatibility for Pillow 10+
    if not hasattr(Image, 'ANTIALIAS'):
        Image.ANTIALIAS = Image.Resampling.LANCZOS
    import pypdfium2 as pdfium
    try:
        import torch
    except Exception:
        torch = None
    
    # Add Word document support
    try:
        import docx
        HAS_DOCX = True
    except ImportError:
        HAS_DOCX = False
    
    # For legacy .doc files, we'll need antiword or similar
    try:
        import subprocess
        import shlex
        HAS_LEGACY_DOC_SUPPORT = True
    except ImportError:
        HAS_LEGACY_DOC_SUPPORT = False
    
    # Import Arabic reshaper from our model
    try:
        from arabic_reshaper import reshape
        from ocr_config import get_config, get_cache_dir
        HAS_MODEL = True
    except ImportError:
        # Fallback if model imports fail
        HAS_MODEL = False
        def reshape(text):
            return text
        def get_config():
            return {"languages": ["ar", "en"], "use_gpu": True}
        def get_cache_dir():
            return Path.home() / ".EasyOCR"
    
    HAS_DEPENDENCIES = True
except ImportError as e:
    HAS_DEPENDENCIES = False
    IMPORT_ERROR = str(e)
    HAS_MODEL = False


def reshape_arabic_text(text: str) -> str:
    """
    Reshape Arabic text using our custom model's reshaper
    This ensures proper Arabic text rendering and ligature support
    Also handles normalization of Arabic numbers and special characters
    """
    if not text:
        return text
    
    try:
        # Normalize Arabic numbers to Latin (optional but recommended for consistency)
        # Arabic: ٠١٢٣٤٥٦٧٨٩  Latin: 0123456789
        arabic_to_latin_map = {
            '\u0660': '0',  # ٠
            '\u0661': '1',  # ١
            '\u0662': '2',  # ٢
            '\u0663': '3',  # ٣
            '\u0664': '4',  # ٤
            '\u0665': '5',  # ٥
            '\u0666': '6',  # ٦
            '\u0667': '7',  # ٧
            '\u0668': '8',  # ٨
            '\u0669': '9',  # ٩
        }
        
        for arabic_digit, latin_digit in arabic_to_latin_map.items():
            text = text.replace(arabic_digit, latin_digit)
        
        # Use the reshape function from our model
        reshaped = reshape(text)
        return reshaped
    except Exception as e:
        # If reshaping fails, return original text
        print(f"Warning: Arabic reshaping failed: {e}", file=sys.stderr)
        return text


def preprocess_image(image_array):
    """
    Minimal preprocessing: grayscale only.
    EasyOCR is a deep-learning model with its own internal preprocessing.
    Heavy image manipulation (thresholding, CLAHE) destroys Arabic text features.
    """
    try:
        # Convert to grayscale if needed — that's all EasyOCR needs
        if len(image_array.shape) == 3:
            return cv2.cvtColor(image_array, cv2.COLOR_BGR2GRAY)
        return image_array
    except Exception as e:
        print(f"Warning: Preprocessing failed, using original: {e}", file=sys.stderr)
        return image_array


def denoise_image(image_array):
    """
    Apply optional denoising for cleaner PDF scans
    Used only when text detection is low
    """
    try:
        # Only denoise if image is color, otherwise skip
        if len(image_array.shape) == 3 and image_array.shape[2] == 3:
            # Light bilateral filter to preserve edges
            denoised = cv2.bilateralFilter(image_array, 5, 20, 20)
            return denoised
        return image_array
    except Exception as e:
        print(f"Warning: Denoising failed: {e}", file=sys.stderr)
        return image_array


def process_image_with_easyocr(image_path: str, reader) -> dict:
    """
    Process a single image using EasyOCR with Arabic support.
    Sorts detections top-to-bottom, and right-to-left within each line.
    """
    try:
        image = cv2.imread(image_path)
        if image is None:
            pil_image = Image.open(image_path)
            image = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)

        processed = preprocess_image(image)

        # detail=1 gives bounding boxes for proper RTL sorting
        results = reader.readtext(processed, detail=1, paragraph=False, batch_size=8)

        if not results:
            return {"success": False, "error": "No text detected in image"}

        # Sort by Y then X descending (RTL for Arabic)
        results.sort(key=lambda r: (r[0][0][1], -r[0][0][0]))

        # Calculate adaptive y_threshold from median text height
        heights = []
        for r in results:
            bbox = r[0]
            h = abs(bbox[2][1] - bbox[0][1])  # bottom - top
            if h > 5:
                heights.append(h)
        median_height = np.median(heights) if heights else 20
        y_threshold = max(20, median_height * 0.6)  # adaptive threshold

        # Group words into lines by Y proximity
        lines = []
        current_line = []
        current_y = None

        for bbox, text, confidence in results:
            if not text.strip():
                continue
            y = bbox[0][1]
            if current_y is None or abs(y - current_y) <= y_threshold:
                current_line.append((bbox, text.strip(), confidence))
                current_y = (current_y * len(current_line) + y) / (len(current_line) + 1) if current_y else y
            else:
                # Sort current line right-to-left (by X descending)
                current_line.sort(key=lambda r: -r[0][0][0])
                lines.append(" ".join([w[1] for w in current_line]))
                current_line = [(bbox, text.strip(), confidence)]
                current_y = y

        if current_line:
            current_line.sort(key=lambda r: -r[0][0][0])
            lines.append(" ".join([w[1] for w in current_line]))

        raw_text = "\n".join(lines)

        if not raw_text.strip():
            return {"success": False, "error": "No text detected in image"}

        return {
            "success": True,
            "text": raw_text,
            "raw_text": raw_text,
            "engine": "easyocr",
            "languages": ["ar", "en"],
            "model": "easyocr"
        }

    except Exception as e:
        return {
            "success": False,
            "error": f"Image processing failed: {str(e)}"
        }


def process_docx_with_text_extraction(docx_path: str) -> dict:
    """
    Process DOCX by extracting text directly (no OCR needed)
    """
    try:
        if not HAS_DOCX:
            return {
                "success": False,
                "error": "python-docx not installed. Install with: pip install python-docx"
            }
        
        doc = docx.Document(docx_path)
        all_paragraphs = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                all_paragraphs.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    all_paragraphs.append(" | ".join(row_text))
        
        combined_text = "\n".join(all_paragraphs)
        
        if not combined_text.strip():
            return {
                "success": False,
                "error": "No text detected in DOCX"
            }
        
        return {
            "success": True,
            "text": combined_text,
            "raw_text": combined_text,
            "engine": "direct_extraction",
            "languages": ["ar", "en"],
            "model": "docx_extraction"
        }
    
    except Exception as e:
        return {
            "success": False,
            "error": f"DOCX processing failed: {str(e)}"
        }


def process_legacy_doc_with_antiword(doc_path: str) -> dict:
    """
    Process legacy .doc files using antiword (if available)
    """
    try:
        if not HAS_LEGACY_DOC_SUPPORT:
            return {
                "success": False,
                "error": "Legacy .doc support not available. Install antiword or convert to .docx"
            }
        
        # Try to use antiword command
        try:
            result = subprocess.run(
                ['antiword', doc_path],
                capture_output=True,
                text=True,
                timeout=30
            )
            
            if result.returncode == 0:
                text = result.stdout
                reshaped_text = reshape_arabic_text(text)

                if not reshaped_text.strip():
                    return {
                        "success": False,
                        "error": "No text detected in legacy DOC"
                    }
                
                return {
                    "success": True,
                    "text": reshaped_text,
                    "raw_text": text,
                    "engine": "antiword_extraction",
                    "languages": ["ar", "en"],
                    "model": "legacy_doc_extraction"
                }
            else:
                return {
                    "success": False,
                    "error": f"antiword failed: {result.stderr}"
                }
                
        except (subprocess.TimeoutExpired, FileNotFoundError) as e:
            return {
                "success": False,
                "error": f"antiword not available or failed: {str(e)}. Please convert .doc to .docx format"
            }
    
    except Exception as e:
        return {
            "success": False,
            "error": f"Legacy .doc processing failed: {str(e)}"
        }


def process_pdf_with_easyocr(pdf_path: str, reader) -> dict:
    """
    Process PDF by converting pages to images and running OCR
    Optimized for large files and Arabic text with adaptive DPI
    """
    try:
        pdf = pdfium.PdfDocument(pdf_path)
        total_pages = len(pdf)
        all_text = []
        pages_processed = 0
        
        # Get file size for adaptive processing
        file_size = os.path.getsize(pdf_path) / (1024 * 1024)  # MB
        
        # Adaptive DPI based on file size, page count, AND language
        # Arabic PDFs need more detail but we compensate with smarter detection
        try:
            cfg = get_config()
            base_dpi = cfg.get("pdf_config", {}).get("dpi", 150)  # Default 150
        except Exception:
            base_dpi = 150
        
        # More aggressive DPI reduction for better speed
        if file_size > 20:  # Large files > 20MB
            dpi = 100
        elif total_pages > 50:  # Many pages
            dpi = 120
        elif total_pages > 20:
            dpi = 140
        else:
            dpi = 150  # Standard files
        
        print(f"Processing PDF: {total_pages} pages, {file_size:.1f}MB, DPI={dpi}", file=sys.stderr)
        
        for page_num in range(total_pages):
            try:
                # Progress reporting
                if (page_num + 1) % 5 == 0 or page_num == total_pages - 1 or page_num == 0:
                    progress = ((page_num + 1) / total_pages) * 100
                    print(f"[OCR] Page {page_num + 1}/{total_pages} ({progress:.0f}%)", file=sys.stderr)
                
                page = pdf[page_num]
                
                # Render page to image
                scale = float(dpi) / 72.0
                bitmap = page.render(scale=scale)
                pil_image = bitmap.to_pil()
                
                # Convert to grayscale numpy array immediately to save memory
                if pil_image.mode != 'L':
                    pil_image = pil_image.convert('L')
                image_array = np.array(pil_image)
                
                # Clean up PIL image and bitmap as soon as possible
                del pil_image
                bitmap.close()
                page.close()
                
                # Light preprocessing - grayscale only for Arabic text
                processed = preprocess_image(image_array)

                # For Arabic text, use paragraph=True for better context
                # detail=1 gives bounding boxes for proper RTL sorting
                results = reader.readtext(
                    processed, 
                    detail=1, 
                    paragraph=False,  # Handle paragraphs manually for better RTL
                    batch_size=8,
                    text_threshold=0.5,  # Lower threshold for faint Arabic text
                    low_text=0.3,  # Lower threshold to catch more text
                )

                if not results:
                    # If no text found, try with more aggressive parameters
                    results = reader.readtext(
                        processed, 
                        detail=1, 
                        batch_size=8,
                        text_threshold=0.3,
                        low_text=0.2,
                    )

                if results:
                    # Sort by Y then X descending (RTL for Arabic)
                    results.sort(key=lambda r: (r[0][0][1], -r[0][0][0]))

                    # Calculate adaptive y_threshold from median text height
                    heights = []
                    for r in results:
                        bbox = r[0]
                        h = abs(bbox[2][1] - bbox[0][1])
                        if h > 3:  # Lower threshold for small text
                            heights.append(h)
                    
                    median_height = np.median(heights) if heights else 15
                    y_threshold = max(10, median_height * 0.5)  # More aggressive grouping

                    # Group words into lines by Y proximity
                    lines = []
                    current_line = []
                    current_y = None

                    for bbox, text, confidence in results:
                        text = text.strip()
                        if not text or confidence < 0.1:  # Skip very low confidence
                            continue
                        
                        y = bbox[0][1]
                        if current_y is None or abs(y - current_y) <= y_threshold:
                            current_line.append((bbox, text, confidence))
                            current_y = (current_y * len(current_line) + y) / (len(current_line) + 1) if current_y else y
                        else:
                            # Sort current line right-to-left (by X descending)
                            current_line.sort(key=lambda r: -r[0][0][0])
                            lines.append(" ".join([w[1] for w in current_line]))
                            current_line = [(bbox, text, confidence)]
                            current_y = y

                    if current_line:
                        current_line.sort(key=lambda r: -r[0][0][0])
                        lines.append(" ".join([w[1] for w in current_line]))

                    page_text = "\n".join(lines)
                else:
                    page_text = ""
                
                if page_text.strip():
                    # Reshape Arabic text for proper rendering
                    reshaped_page = reshape_arabic_text(page_text)
                    all_text.append(f"--- Page {page_num + 1} ---\n{reshaped_page}")
                    pages_processed += 1
                
                # Explicit garbage collection every 3 pages to prevent memory spikes
                if (page_num + 1) % 3 == 0:
                    gc.collect()
                
            except Exception as page_error:
                print(f"Warning: Failed to process page {page_num + 1}: {page_error}", file=sys.stderr)
                continue
        
        combined_text = "\n\n".join(all_text)
        pdf.close()  # Close document
        gc.collect()

        if not combined_text.strip():
            return {
                "success": False,
                "error": "No text detected in PDF",
                "details": f"Processed {total_pages} pages with DPI={dpi}"
            }

        return {
            "success": True,
            "text": combined_text,
            "raw_text": combined_text,
            "engine": "easyocr",
            "languages": ["ar", "en"],
            "pages_processed": pages_processed,
            "total_pages": total_pages,
            "dpi_used": dpi,
            "file_size_mb": round(file_size, 2),
            "model": "easyocr"
        }
    
    except Exception as e:
        return {
            "success": False,
            "error": f"PDF processing failed: {str(e)}",
            "traceback": str(e)
        }


def main():
    """
    Main entry point for OCR processing
    Expects file path as command line argument
    Returns JSON result to stdout
    """
    if not HAS_DEPENDENCIES:
        result = {
            "success": False,
            "error": f"Missing required dependencies: {IMPORT_ERROR}",
            "hint": "Install: pip install easyocr opencv-python-headless pypdfium2 pillow"
        }
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(1)
    
    if len(sys.argv) < 2:
        result = {
            "success": False,
            "error": "No file path provided",
            "usage": "python ocr_runner.py <file_path>"
        }
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    if not os.path.exists(file_path):
        result = {
            "success": False,
            "error": f"File not found: {file_path}"
        }
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(1)
    
    # Load configuration
    config = get_config()
    
    # Determine file type
    file_ext = os.path.splitext(file_path)[1].lower()
    
    try:
        # Initialize EasyOCR reader with configuration
        languages = config["languages"]
        use_gpu = config["use_gpu"]
        
        print(f"Initializing EasyOCR with languages: {languages}, GPU: {use_gpu}...", file=sys.stderr)
        try:
            if torch is not None:
                print(f"torch: {torch.__version__}", file=sys.stderr)
                print(f"torch.cuda.is_available: {torch.cuda.is_available()}", file=sys.stderr)
            else:
                print("torch: not installed", file=sys.stderr)
        except Exception:
            pass
        
        try:
            # Try GPU first if enabled
            if use_gpu:
                reader = easyocr.Reader(
                    languages, 
                    gpu=True, 
                    verbose=False,
                    model_storage_directory=str(get_cache_dir())
                )
                device = "cuda"
            else:
                raise Exception("GPU disabled in config")
        except Exception:
            # Fallback to CPU
            reader = easyocr.Reader(
                languages, 
                gpu=False, 
                verbose=False,
                model_storage_directory=str(get_cache_dir())
            )
            device = "cpu"
        
        print(f"EasyOCR initialized on {device}", file=sys.stderr)
        
        # Process based on file type
        if file_ext == '.pdf':
            result = process_pdf_with_easyocr(file_path, reader)
        elif file_ext == '.docx':
            # DOCX files don't need OCR - direct text extraction
            result = process_docx_with_text_extraction(file_path)
        elif file_ext == '.doc':
            # Legacy DOC files need antiword or conversion
            result = process_legacy_doc_with_antiword(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.webp']:
            result = process_image_with_easyocr(file_path, reader)
        else:
            result = {
                "success": False,
                "error": f"Unsupported file type: {file_ext}",
                "supported": [".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"]
            }
        
        # Add device info and config to result
        if result.get("success"):
            result["device"] = device
            result["config"] = {
                "languages": languages,
                "gpu_enabled": use_gpu,
                "gpu_used": device == "cuda"
            }
        
        # Output JSON result
        print(json.dumps(result, ensure_ascii=False))
        
        sys.exit(0 if result.get("success") else 1)
    
    except Exception as e:
        result = {
            "success": False,
            "error": f"OCR processing failed: {str(e)}",
            "traceback": str(e)
        }
        print(json.dumps(result, ensure_ascii=False))
        sys.exit(1)


if __name__ == "__main__":
    main()
