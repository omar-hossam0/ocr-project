"""
Camera / Image/PDF OCR → PDF/TXT/DOCX
─────────────────────────
At startup a menu appears — choose:
  1 → Camera live capture
  2 → Load image(s) from disk
    3 → Load PDF from disk
"""

import cv2
import numpy as np
import tkinter as tk
from tkinter import filedialog, messagebox
import easyocr
import torch
import pypdfium2 as pdfium
from fpdf import FPDF
import datetime
import os
import tempfile
from pathlib import Path
import warnings
import re
from urllib.parse import quote_plus

from PIL import Image
from transformers import TrOCRProcessor, VisionEncoderDecoderModel
import arabic_reshaper
from bidi.algorithm import get_display

try:
    from pyzbar.pyzbar import decode as pyzbar_decode
except Exception:
    pyzbar_decode = None

try:
    import zxingcpp
except Exception:
    zxingcpp = None

ORIENTATION_LABELS = {0: "0°", 1: "90°", 2: "180°", 3: "270°"}

# Quiet known EasyOCR numeric warnings that do not affect final text output.
warnings.filterwarnings("ignore", message="overflow encountered in scalar add", category=RuntimeWarning)

trocr_processor = None
trocr_model = None
trocr_device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
easyocr_reader = None
_trocr_checked_once = False
_easyocr_loaded_once = False
ENABLE_TROCR = os.environ.get("ENABLE_TROCR", "0") == "1"
qr_detector = cv2.QRCodeDetector()
QR_RESOLVER_TEMPLATE = os.environ.get("QR_RESOLVER_TEMPLATE", "").strip()


def _get_easyocr_reader():
    """Lazy-load EasyOCR reader so UI starts instantly."""
    global easyocr_reader, _easyocr_loaded_once
    if easyocr_reader is None:
        if not _easyocr_loaded_once:
            print("Loading OCR model (EasyOCR Arabic+English, first run downloads weights) ...")
        easyocr_reader = easyocr.Reader(["ar", "en"], gpu=torch.cuda.is_available())
        _easyocr_loaded_once = True
        print("EasyOCR ready.\n")
    return easyocr_reader


def _init_optional_trocr():
    """Load local fine-tuned TrOCR model if available."""
    global trocr_processor, trocr_model, _trocr_checked_once
    if _trocr_checked_once:
        return
    _trocr_checked_once = True

    # Default is fast mode: do not load TrOCR unless explicitly enabled.
    if not ENABLE_TROCR:
        print("Fast mode: TrOCR disabled (set ENABLE_TROCR=1 to enable).")
        return

    # TrOCR on CPU is slow for interactive OCR; keep it GPU-only by default.
    if not torch.cuda.is_available():
        print("Fast mode: TrOCR skipped on CPU. Use EasyOCR only.")
        return

    model_dir = Path(__file__).resolve().parent / "arabic_trocr_model"
    if not model_dir.exists():
        print("No local arabic_trocr_model found, using EasyOCR recognizer.")
        return
    try:
        print(f"Loading local TrOCR model: {model_dir}")
        trocr_processor = TrOCRProcessor.from_pretrained(str(model_dir))
        trocr_model = VisionEncoderDecoderModel.from_pretrained(str(model_dir)).to(trocr_device)
        trocr_model.eval()
        print(f"TrOCR loaded on {trocr_device}.\n")
    except Exception as e:
        print(f"[WARN] Failed to load arabic_trocr_model: {e}")
        trocr_processor = None
        trocr_model = None

# ══════════════════════════════════════════════════════════════
# STARTUP MENU — choose mode before anything opens
# ══════════════════════════════════════════════════════════════
def _show_menu():
    """Show a simple Tk menu and return 'camera', 'image' or 'pdf'."""
    choice = {"value": None}

    root = tk.Tk()
    root.title("OCR → PDF")
    root.resizable(False, False)
    root.attributes("-topmost", True)

    # Centre on screen
    root.update_idletasks()
    w, h = 460, 220
    x = (root.winfo_screenwidth()  - w) // 2
    y = (root.winfo_screenheight() - h) // 2
    root.geometry(f"{w}x{h}+{x}+{y}")

    tk.Label(root, text="OCR → PDF", font=("Helvetica", 16, "bold"), pady=12).pack()
    tk.Label(root, text="اختر طريقة الإدخال:", font=("Helvetica", 11)).pack()

    btn_frame = tk.Frame(root, pady=14)
    btn_frame.pack()

    def pick_camera():
        choice["value"] = "camera"
        root.destroy()

    def pick_image():
        choice["value"] = "image"
        root.destroy()

    def pick_pdf():
        choice["value"] = "pdf"
        root.destroy()

    tk.Button(btn_frame, text="📷  Camera", width=12, height=2,
              font=("Helvetica", 11), bg="#27ae60", fg="white",
              command=pick_camera).grid(row=0, column=0, padx=10)

    tk.Button(btn_frame, text="🖼️  Import Picture ", width=18, height=2,
              font=("Helvetica", 11), bg="#2980b9", fg="white",
              command=pick_image).grid(row=0, column=1, padx=10)

    tk.Button(btn_frame, text="📄  Import PDF", width=14, height=2,
              font=("Helvetica", 11), bg="#8e44ad", fg="white",
              command=pick_pdf).grid(row=0, column=2, padx=10)

    root.protocol("WM_DELETE_WINDOW", root.destroy)
    root.mainloop()
    return choice["value"]

MODE = _show_menu()
if MODE is None:
    print("Cancelled.")
    exit(0)

# ── Shared helpers ────────────────────────────────────────────
def _unpack_ori(pred):
    if isinstance(pred, (list, tuple)) and len(pred) == 1:
        pred = pred[0]
    if isinstance(pred, (list, tuple)) and len(pred) >= 2:
        return int(pred[0]), float(pred[1])
    return int(pred), 1.0


def _is_arabic_char(ch):
    code = ord(ch)
    return (
        0x0600 <= code <= 0x06FF
        or 0x0750 <= code <= 0x077F
        or 0x08A0 <= code <= 0x08FF
        or 0xFB50 <= code <= 0xFDFF
        or 0xFE70 <= code <= 0xFEFF
    )


def _shape_arabic_for_pdf(text):
    """Reshape Arabic text and apply bidi for proper PDF rendering."""
    if not text:
        return text
    if not any(_is_arabic_char(ch) for ch in text):
        return text
    try:
        reshaped = arabic_reshaper.reshape(text)
        return get_display(reshaped)
    except Exception:
        return text


def _arabic_ratio(text):
    chars = [c for c in text if c.strip()]
    if not chars:
        return 0.0
    ar = sum(1 for c in chars if _is_arabic_char(c))
    return ar / len(chars)


def _dedupe_keep_order(items):
    seen = set()
    out = []
    for item in items:
        if item not in seen:
            seen.add(item)
            out.append(item)
    return out


def _normalize_qr_payload(payload):
    text = str(payload).strip()
    if not text:
        return ""

    for prefix in ("QRCODE:", "QR Code:", "QR:", "CODE128:", "CODE39:"):
        if text.upper().startswith(prefix.upper()):
            text = text[len(prefix):].strip()
            break

    return text


def _is_url(text):
    return bool(re.match(r"^https?://", text.strip(), flags=re.IGNORECASE))


def _build_qr_links(qr_payloads):
    """Build best-effort links from QR payloads.

    If payload already contains URL → use it.
    If QR_RESOLVER_TEMPLATE is set → format with {code}.
    Always add a search link fallback for non-URL payloads.
    """
    links = []
    for payload in qr_payloads:
        value = _normalize_qr_payload(payload)
        if not value:
            continue

        if _is_url(value):
            links.append(value)
            continue

        if QR_RESOLVER_TEMPLATE and "{code}" in QR_RESOLVER_TEMPLATE:
            links.append(QR_RESOLVER_TEMPLATE.format(code=value))

        links.append(f"https://www.google.com/search?q={quote_plus(value)}")

    return _dedupe_keep_order(links)


def _extract_qr_payloads(image_rgb):
    """Extract QR payloads using OpenCV + optional pyzbar fallback."""
    payloads = []

    def _decode_opencv(img):
        out = []
        try:
            retval, decoded_info, _points, _ = qr_detector.detectAndDecodeMulti(img)
            if retval and decoded_info:
                out.extend([str(x).strip() for x in decoded_info if str(x).strip()])
        except Exception:
            pass

        try:
            single, _points, _ = qr_detector.detectAndDecode(img)
            single = str(single).strip()
            if single:
                out.append(single)
        except Exception:
            pass

        if hasattr(qr_detector, "detectAndDecodeCurved"):
            try:
                curved, _points, _ = qr_detector.detectAndDecodeCurved(img)
                curved = str(curved).strip()
                if curved:
                    out.append(curved)
            except Exception:
                pass

        return out

    gray = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2GRAY)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8)).apply(gray)
    _, otsu = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    upscaled = cv2.resize(gray, None, fx=2.2, fy=2.2, interpolation=cv2.INTER_CUBIC)
    blur = cv2.GaussianBlur(gray, (3, 3), 0)

    variants = [
        image_rgb,
        gray,
        clahe,
        otsu,
        blur,
        upscaled,
    ]

    for variant in variants:
        payloads.extend(_decode_opencv(variant))

    if pyzbar_decode is not None:
        for variant in variants:
            try:
                decoded = pyzbar_decode(variant)
                for item in decoded:
                    data = item.data.decode("utf-8", errors="ignore").strip()
                    if data:
                        kind = getattr(item, "type", "CODE")
                        payloads.append(f"{kind}: {data}")
            except Exception:
                pass

    # ZXing is often stronger on rotated, low-contrast, or non-QR symbologies.
    if zxingcpp is not None:
        for variant in variants:
            try:
                results = zxingcpp.read_barcodes(variant)
                for res in results:
                    text = str(getattr(res, "text", "")).strip()
                    if text:
                        fmt = str(getattr(res, "format", "CODE"))
                        payloads.append(f"{fmt}: {text}")
            except Exception:
                pass

    # If we have detected QR corners, try rectifying the perspective and decode again.
    try:
        found, points = qr_detector.detect(gray)
        if found and points is not None and len(points) > 0:
            pts = points[0].astype(np.float32)
            side = int(max(np.linalg.norm(pts[0] - pts[1]), np.linalg.norm(pts[1] - pts[2]), 120))
            dst = np.array([[0, 0], [side - 1, 0], [side - 1, side - 1], [0, side - 1]], dtype=np.float32)
            m = cv2.getPerspectiveTransform(pts, dst)
            warped = cv2.warpPerspective(gray, m, (side, side))
            variants2 = [warped, cv2.resize(warped, None, fx=2.0, fy=2.0, interpolation=cv2.INTER_CUBIC)]

            for v in variants2:
                payloads.extend(_decode_opencv(v))
                if pyzbar_decode is not None:
                    try:
                        decoded = pyzbar_decode(v)
                        for item in decoded:
                            data = item.data.decode("utf-8", errors="ignore").strip()
                            if data:
                                kind = getattr(item, "type", "CODE")
                                payloads.append(f"{kind}: {data}")
                    except Exception:
                        pass
                if zxingcpp is not None:
                    try:
                        results = zxingcpp.read_barcodes(v)
                        for res in results:
                            text = str(getattr(res, "text", "")).strip()
                            if text:
                                fmt = str(getattr(res, "format", "CODE"))
                                payloads.append(f"{fmt}: {text}")
                    except Exception:
                        pass
    except Exception:
        pass

    normalized = [_normalize_qr_payload(p) for p in payloads]
    normalized = [p for p in normalized if p]
    return _dedupe_keep_order(normalized)


def _merge_ocr_and_qr(ocr_text, qr_payloads):
    if not qr_payloads:
        return (ocr_text or "").strip()

    urls = _build_qr_links(qr_payloads)

    qr_lines = ["[QR DETECTED]"]
    if urls:
        qr_lines.append("[QR LINKS]")
        qr_lines.extend([f"- {u}" for u in _dedupe_keep_order(urls)])
    qr_lines.extend([f"- {q}" for q in qr_payloads])

    if (ocr_text or "").strip():
        return "\n".join(qr_lines + ["", ocr_text.strip()])
    return "\n".join(qr_lines)


def _prepare_ocr_variants(image_rgb):
    """Generate a small set of image variants to improve OCR robustness."""
    variants = []
    variants.append(("rgb", image_rgb))

    gray = cv2.cvtColor(image_rgb, cv2.COLOR_RGB2GRAY)
    variants.append(("gray", gray))

    den = cv2.bilateralFilter(gray, 9, 75, 75)
    th = cv2.adaptiveThreshold(
        den,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        31,
        15,
    )
    variants.append(("adaptive", th))

    h, w = gray.shape[:2]
    scale = 2.0 if max(h, w) < 1800 else 1.5
    # Keep CPU path fast by limiting to two variants.
    if torch.cuda.is_available():
        up = cv2.resize(gray, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_CUBIC)
        variants.append(("upscaled", up))

    return variants


def _read_easyocr(image_variant):
    """Return text and confidence using EasyOCR detail output."""
    reader = _get_easyocr_reader()
    _init_optional_trocr()

    detections = reader.readtext(
        image_variant,
        detail=1,
        paragraph=False,
        decoder="greedy",
        beamWidth=1,
        contrast_ths=0.05,
        adjust_contrast=0.7,
        text_threshold=0.65,
        low_text=0.25,
        link_threshold=0.4,
    )

    if not detections:
        return "", 0.0

    # Sort roughly top-to-bottom then right-to-left for Arabic-heavy lines.
    ordered = sorted(
        detections,
        key=lambda d: (
            (d[0][0][1] + d[0][2][1]) / 2.0,
            -((d[0][0][0] + d[0][2][0]) / 2.0),
        ),
    )

    lines = []
    confs = []
    for box, txt, conf in ordered:
        text = str(txt).strip()

        # If a local Arabic TrOCR exists, use it as recognizer per detected box.
        if trocr_model is not None and trocr_processor is not None:
            try:
                xs = [int(p[0]) for p in box]
                ys = [int(p[1]) for p in box]
                x1, x2 = max(0, min(xs)), min(image_variant.shape[1], max(xs))
                y1, y2 = max(0, min(ys)), min(image_variant.shape[0], max(ys))

                # Add a small padding around detected box.
                pad_x = max(2, int((x2 - x1) * 0.06))
                pad_y = max(2, int((y2 - y1) * 0.12))
                x1 = max(0, x1 - pad_x)
                y1 = max(0, y1 - pad_y)
                x2 = min(image_variant.shape[1], x2 + pad_x)
                y2 = min(image_variant.shape[0], y2 + pad_y)

                if x2 > x1 and y2 > y1:
                    crop = image_variant[y1:y2, x1:x2]
                    if len(crop.shape) == 2:
                        crop_rgb = cv2.cvtColor(crop, cv2.COLOR_GRAY2RGB)
                    else:
                        crop_rgb = crop

                    pil_crop = Image.fromarray(crop_rgb)
                    pixel_values = trocr_processor(images=pil_crop, return_tensors="pt").pixel_values.to(trocr_device)
                    with torch.no_grad():
                        generated = trocr_model.generate(pixel_values, max_length=128)
                    trocr_text = trocr_processor.batch_decode(generated, skip_special_tokens=True)[0].strip()

                    if trocr_text and len(trocr_text) >= max(2, len(text) // 3):
                        text = trocr_text
            except Exception:
                pass

        if text:
            lines.append(text)
            confs.append(float(conf))

    joined = "\n".join(lines).strip()
    if not joined:
        return "", 0.0

    avg_conf = sum(confs) / len(confs) if confs else 0.0
    return joined, avg_conf


def _ocr_text_from_image(image_rgb):
    """Read text using multiple preprocessed variants and keep best candidate."""
    candidates = []
    for variant_name, variant_image in _prepare_ocr_variants(image_rgb):
        text, avg_conf = _read_easyocr(variant_image)
        if text:
            ratio = _arabic_ratio(text)
            # Prefer higher confidence and stronger Arabic coverage for Arabic docs.
            score = (avg_conf * 0.8) + (ratio * 0.2)
            candidates.append((score, avg_conf, ratio, variant_name, text))

    if not candidates:
        return ""

    best = max(candidates, key=lambda c: c[0])
    print(
        f"  OCR variant: {best[3]} | conf={best[1]:.3f} | ar_ratio={best[2]:.3f}"
    )
    return best[4]

def run_ocr(img_path):
    """Run full OCR pipeline on an image file. Returns text string."""
    bgr = cv2.imread(img_path)
    if bgr is None:
        raise ValueError(f"Cannot read image: {img_path}")
    rgb = cv2.cvtColor(bgr, cv2.COLOR_BGR2RGB)
    qr_payloads = _extract_qr_payloads(rgb)
    if qr_payloads:
        print("  QR found:")
        for q in qr_payloads:
            print(f"    {q}")
        links = _build_qr_links(qr_payloads)
        if links:
            print("  QR links:")
            for u in _dedupe_keep_order(links):
                print(f"    {u}")
    ocr_text = _ocr_text_from_image(rgb)
    return _merge_ocr_and_qr(ocr_text, qr_payloads)

def save_pdf(pages_text):
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    pdf_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        f"ocr_result_{timestamp}.pdf"
    )
    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)

    font_name = "Helvetica"
    unicode_font = None
    for candidate in [
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\tahoma.ttf",
        r"C:\Windows\Fonts\times.ttf",
    ]:
        if os.path.exists(candidate):
            unicode_font = candidate
            break

    if unicode_font:
        try:
            pdf.add_font("Unicode", "", unicode_font, uni=True)
            font_name = "Unicode"
        except Exception:
            font_name = "Helvetica"

    # A4 width=210mm, margins 15mm each side → usable = 180mm
    TEXT_W = 180
    for page_num, page_text in enumerate(pages_text, 1):
        pdf.add_page()
        pdf.set_left_margin(15)
        pdf.set_right_margin(15)
        pdf.set_font(font_name, style="", size=14)
        pdf.set_x(15)
        title = _shape_arabic_for_pdf(f"Capture {page_num}")
        pdf.cell(TEXT_W, 10, title, new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(font_name, size=11)
        pdf.ln(2)
        for line in page_text.splitlines():
            pdf.set_x(15)
            if line.strip():
                has_arabic = any(_is_arabic_char(ch) for ch in line)
                shaped_line = _shape_arabic_for_pdf(line)
                if font_name == "Helvetica":
                    safe = shaped_line.encode("latin-1", errors="replace").decode("latin-1")
                    pdf.multi_cell(TEXT_W, 7, safe, align="R" if has_arabic else "L")
                else:
                    pdf.multi_cell(TEXT_W, 7, shaped_line, align="R" if has_arabic else "L")
            else:
                pdf.ln(4)
        pdf.ln(4)
    pdf.output(pdf_path)
    return pdf_path


def save_txt(pages_text):
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    txt_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        f"ocr_result_{timestamp}.txt"
    )

    with open(txt_path, "w", encoding="utf-8") as f:
        for page_num, page_text in enumerate(pages_text, 1):
            f.write(f"===== Page {page_num} =====\n")
            f.write(page_text)
            f.write("\n\n")

    return txt_path


def save_docx(pages_text):
    try:
        from docx import Document
    except Exception as e:
        raise RuntimeError(
            "python-docx is not installed. Install it with: pip install python-docx"
        ) from e

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    docx_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        f"ocr_result_{timestamp}.docx"
    )

    doc = Document()
    doc.add_heading("OCR Result", level=1)
    for page_num, page_text in enumerate(pages_text, 1):
        doc.add_heading(f"Page {page_num}", level=2)
        doc.add_paragraph(page_text if page_text.strip() else "(No text detected)")

    doc.save(docx_path)
    return docx_path


def ask_output_format():
    """Ask user which output format they want."""
    choice = {"value": None}

    root = tk.Tk()
    root.title("OCR Output")
    root.resizable(False, False)
    root.attributes("-topmost", True)

    root.update_idletasks()
    w, h = 500, 190
    x = (root.winfo_screenwidth() - w) // 2
    y = (root.winfo_screenheight() - h) // 2
    root.geometry(f"{w}x{h}+{x}+{y}")

    tk.Label(root, text="Choose output format", font=("Helvetica", 14, "bold"), pady=12).pack()
    tk.Label(root, text="اختر صيغة الحفظ:", font=("Helvetica", 11)).pack()

    btn_frame = tk.Frame(root, pady=16)
    btn_frame.pack()

    def pick(fmt):
        choice["value"] = fmt
        root.destroy()

    tk.Button(
        btn_frame, text="PDF", width=12, height=2,
        font=("Helvetica", 11), bg="#2980b9", fg="white",
        command=lambda: pick("pdf")
    ).grid(row=0, column=0, padx=8)

    tk.Button(
        btn_frame, text="TXT", width=12, height=2,
        font=("Helvetica", 11), bg="#27ae60", fg="white",
        command=lambda: pick("txt")
    ).grid(row=0, column=1, padx=8)

    tk.Button(
        btn_frame, text="DOCX", width=12, height=2,
        font=("Helvetica", 11), bg="#8e44ad", fg="white",
        command=lambda: pick("docx")
    ).grid(row=0, column=2, padx=8)

    root.protocol("WM_DELETE_WINDOW", root.destroy)
    root.mainloop()
    return choice["value"]


def save_output(pages_text):
    fmt = ask_output_format()
    if fmt is None:
        print("Save cancelled.")
        return None

    if fmt == "pdf":
        return save_pdf(pages_text)
    if fmt == "txt":
        return save_txt(pages_text)
    if fmt == "docx":
        return save_docx(pages_text)

    return None


def run_ocr_pdf(pdf_path):
    """Run OCR on all pages of a PDF and return list of page texts."""
    pdf = pdfium.PdfDocument(pdf_path)
    pages = []
    for page_idx in range(len(pdf)):
        print(f"  OCR page {page_idx + 1} ...")
        page = pdf[page_idx]
        bitmap = page.render(scale=2.0)
        pil_img = bitmap.to_pil()
        page_np = np.array(pil_img.convert("RGB"))
        qr_payloads = _extract_qr_payloads(page_np)
        if qr_payloads:
            print("  QR found on page:")
            for q in qr_payloads:
                print(f"    {q}")
            links = _build_qr_links(qr_payloads)
            if links:
                print("  QR links on page:")
                for u in _dedupe_keep_order(links):
                    print(f"    {u}")
        page_text = _ocr_text_from_image(page_np)
        page_text = _merge_ocr_and_qr(page_text, qr_payloads)
        pages.append(page_text)
    return pages

all_pages_text = []

# ══════════════════════════════════════════════════════════════
# MODE A — Image from disk
# ══════════════════════════════════════════════════════════════
if MODE == "image":
    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    paths = filedialog.askopenfilenames(
        title="اختر صورة أو أكثر",
        filetypes=[("Images", "*.jpg *.jpeg *.png *.bmp *.tiff *.tif"), ("All files", "*.*")]
    )
    root.destroy()

    if not paths:
        print("No files selected.")
        exit(0)

    for p in paths:
        print(f"\nProcessing: {p}")
        try:
            text = run_ocr(p)
            if text:
                all_pages_text.append(text)
                print(f"--- text ---\n{text}\n------------")
            else:
                print("  (no text detected)")
        except Exception as e:
            print(f"  [ERROR] {e}")

    if all_pages_text:
        try:
            out_path = save_output(all_pages_text)
        except Exception as e:
            print(f"[ERROR] Save failed: {e}")
            out_path = None

        if out_path:
            print(f"\n[SAVED] Output → {out_path}")
            root2 = tk.Tk(); root2.withdraw()
            messagebox.showinfo("Done", f"Output saved:\n{out_path}")
            root2.destroy()
    else:
        print("Nothing to save.")
    exit(0)

# ══════════════════════════════════════════════════════════════
# MODE B — Camera
# ══════════════════════════════════════════════════════════════
if MODE == "pdf":
    root = tk.Tk()
    root.withdraw()
    root.attributes("-topmost", True)
    pdf_path = filedialog.askopenfilename(
        title="اختر ملف PDF",
        filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
    )
    root.destroy()

    if not pdf_path:
        print("No PDF selected.")
        exit(0)

    print(f"\nProcessing PDF: {pdf_path}")
    try:
        all_pages_text = run_ocr_pdf(pdf_path)
    except Exception as e:
        print(f"[ERROR] PDF OCR failed: {e}")
        exit(1)

    non_empty = [t for t in all_pages_text if t.strip()]
    if not non_empty:
        print("No text detected in PDF pages.")
        exit(0)

    try:
        out_path = save_output(all_pages_text)
    except Exception as e:
        print(f"[ERROR] Save failed: {e}")
        out_path = None

    if out_path:
        print(f"\n[SAVED] OCR output → {out_path}")
        root2 = tk.Tk(); root2.withdraw()
        messagebox.showinfo("Done", f"OCR output saved:\n{out_path}")
        root2.destroy()
    exit(0)

print("Press  SPACE → capture & OCR   |   S → save PDF   |   Q → quit\n")

# ── Open webcam ───────────────────────────────────────────────
cap = cv2.VideoCapture(0)
if not cap.isOpened():
    print("[ERROR] Cannot open camera.")
    exit(1)

# Set camera to higher resolution
cap.set(cv2.CAP_PROP_FRAME_WIDTH, 1280)
cap.set(cv2.CAP_PROP_FRAME_HEIGHT, 720)

# Read one frame just to know the actual resolution
ret0, frame0 = cap.read()
fh, fw = frame0.shape[:2]

# Nice landscape window
cv2.namedWindow("Camera OCR", cv2.WINDOW_NORMAL)
cv2.resizeWindow("Camera OCR", 900, 560)

# A4-proportioned portrait rectangle centred in the frame
# A4 ratio = 1 : 1.414  (width : height)
rect_h = int(fh * 0.96)                  # 96% of frame height
rect_w = int(rect_h / 1.414)             # A4 width from height
MARGIN_X = (fw - rect_w) // 2
MARGIN_Y = (fh - rect_h) // 2
ROI = (MARGIN_X, MARGIN_Y, fw - MARGIN_X, fh - MARGIN_Y)   # x1,y1,x2,y2

while True:
    ret, frame = cap.read()
    if not ret:
        print("[ERROR] Failed to read frame.")
        break

    x1, y1, x2, y2 = ROI

    # ── Draw guide rectangle ──────────────────────────────────
    display = frame.copy()

    # Dim area outside the rectangle
    overlay = display.copy()
    cv2.rectangle(overlay, (0, 0), (fw, fh), (0, 0, 0), -1)
    cv2.rectangle(overlay, (x1, y1), (x2, y2), (0, 0, 0), -1)
    alpha = 0.35
    cv2.addWeighted(overlay, alpha, display, 1 - alpha, 0, display)

    # Draw the guide rectangle itself
    color = (0, 220, 0)
    cv2.rectangle(display, (x1, y1), (x2, y2), color, 3)

    # Corner accents
    corner_len = 25
    thickness  = 4
    for cx, cy, dx, dy in [
        (x1, y1,  1,  1), (x2, y1, -1,  1),
        (x1, y2,  1, -1), (x2, y2, -1, -1),
    ]:
        cv2.line(display, (cx, cy), (cx + dx * corner_len, cy), color, thickness)
        cv2.line(display, (cx, cy), (cx, cy + dy * corner_len), color, thickness)

    # Counter badge
    badge = f"Captures: {len(all_pages_text)}"
    cv2.rectangle(display, (fw - 195, 5), (fw - 5, 35), (0, 0, 0), -1)
    cv2.putText(display, badge, (fw - 190, 27),
                cv2.FONT_HERSHEY_SIMPLEX, 0.65, (0, 220, 0), 2)

    # Instructions
    cv2.putText(display, "Place paper inside the box",
                (x1 + 8, y1 - 10), cv2.FONT_HERSHEY_SIMPLEX, 0.65, color, 2)
    cv2.putText(display, "SPACE: capture | S: save PDF | Q: quit",
                (10, fh - 12), cv2.FONT_HERSHEY_SIMPLEX, 0.6, (255, 255, 255), 2)

    cv2.imshow("Camera OCR", display)
    key = cv2.waitKey(1) & 0xFF

    # ── Capture & OCR ─────────────────────────────────────────
    if key == ord(" "):
        print("Captured! Running OCR ...")

        # Crop only the ROI region
        roi_bgr = frame[y1:y2, x1:x2]

        with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
            tmp_path = tmp.name
        cv2.imwrite(tmp_path, roi_bgr)

        text = run_ocr(tmp_path)
        os.unlink(tmp_path)

        if text:
            all_pages_text.append(text)
            print(f"\n--- Recognised text ---\n{text}\n-----------------------\n")
        else:
            print("  (no text detected)\n")

        # Flash green fill inside ROI to confirm capture
        flash = frame.copy()
        cv2.rectangle(flash, (x1, y1), (x2, y2), (0, 255, 0), -1)
        cv2.addWeighted(flash, 0.3, frame, 0.7, 0, flash)
        cv2.imshow("Camera OCR", flash)
        cv2.waitKey(400)

    # ── Save output ───────────────────────────────────────────
    elif key == ord("s"):
        if not all_pages_text:
            print("[INFO] Nothing captured yet — press SPACE first.")
            continue
        try:
            out_path = save_output(all_pages_text)
            if out_path:
                print(f"[SAVED] Output → {out_path}\n")
        except Exception as e:
            print(f"[ERROR] Save failed: {e}\n")

    # ── Quit ──────────────────────────────────────────────────
    elif key == ord("q"):
        break

cap.release()
cv2.destroyAllWindows()
print("Done.")
